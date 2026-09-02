#!/usr/bin/env python3
"""Build the MoonCen architecture review Word documents from Markdown sources.

The Markdown files and SVG diagrams are the editable source.  This script renders
SVG diagrams to PNG and creates styled DOCX deliverables without changing any
application or deployment source.
"""

from __future__ import annotations

import math
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor
from PIL import Image, ImageDraw, ImageFont

try:
    import cairosvg
except (ImportError, OSError):
    # Minimal CI images often omit the native cairo shared library.  A native
    # Pillow renderer below keeps the Word deliverables reproducible there.
    cairosvg = None


ROOT = Path(__file__).resolve().parent
ASSETS = ROOT / "assets"
SNAPSHOT = "2026-08-19 UTC"
REVISION = "master@8d55e873bfb06ec33f566839fce7ee98650955f8"

SOURCE_FILES = [
    ROOT / "MoonCen_아키텍처_간략본.md",
    ROOT / "MoonCen_아키텍처_상세본.md",
    ROOT / "MoonCen_개선방안_우선순위_간략본.md",
    ROOT / "MoonCen_개선방안_우선순위_상세본.md",
    ROOT / "MoonCen_운영_가이드.md",
]

NAVY = "17324D"
BLUE = "2D6A9F"
LIGHT_BLUE = "EAF2F8"
LIGHT_GRAY = "F2F4F6"
MID_GRAY = "D8DEE5"
TEXT = "273444"
RED = "A94442"
WHITE = "FFFFFF"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=90, start=100, bottom=90, end=100) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for margin, value in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{margin}"))
        if node is None:
            node = OxmlElement(f"w:{margin}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def set_keep_with_next(paragraph, value=True) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    keep = p_pr.find(qn("w:keepNext"))
    if keep is None:
        keep = OxmlElement("w:keepNext")
        p_pr.append(keep)
    keep.set(qn("w:val"), "1" if value else "0")


def set_repeat_table_width(table) -> None:
    tbl_pr = table._tbl.tblPr
    width = tbl_pr.find(qn("w:tblW"))
    if width is None:
        width = OxmlElement("w:tblW")
        tbl_pr.append(width)
    width.set(qn("w:w"), "5000")
    width.set(qn("w:type"), "pct")


def set_run_font(run, name="Malgun Gothic", size=None, color=None, bold=None) -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    if bold is not None:
        run.bold = bold


def setup_document() -> Document:
    doc = Document()
    sec = doc.sections[0]
    sec.page_height = Cm(29.7)
    sec.page_width = Cm(21.0)
    sec.top_margin = Cm(1.8)
    sec.bottom_margin = Cm(1.7)
    sec.left_margin = Cm(1.65)
    sec.right_margin = Cm(1.65)
    sec.header_distance = Cm(0.7)
    sec.footer_distance = Cm(0.7)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Malgun Gothic"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    normal.font.size = Pt(9.6)
    normal.font.color.rgb = RGBColor.from_string(TEXT)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.18
    normal.paragraph_format.space_after = Pt(5)

    for level, size, color, before, after in (
        (1, 23, NAVY, 18, 10),
        (2, 16, NAVY, 15, 7),
        (3, 12.5, BLUE, 11, 5),
        (4, 10.5, TEXT, 8, 4),
    ):
        style = styles[f"Heading {level}"]
        style.font.name = "Malgun Gothic"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    for name in ("List Bullet", "List Bullet 2", "List Bullet 3", "List Number", "List Number 2", "List Number 3"):
        if name in styles:
            style = styles[name]
            style.font.name = "Malgun Gothic"
            style._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
            style.font.size = Pt(9.4)
            style.paragraph_format.space_after = Pt(2.5)

    if "Code Block" not in styles:
        code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles["Code Block"]
    code.font.name = "Consolas"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    code.font.size = Pt(8.2)
    code.font.color.rgb = RGBColor.from_string("243447")
    code.paragraph_format.left_indent = Cm(0.35)
    code.paragraph_format.right_indent = Cm(0.25)
    code.paragraph_format.space_before = Pt(3)
    code.paragraph_format.space_after = Pt(5)
    code.paragraph_format.line_spacing = 1.0

    if "Caption Korean" not in styles:
        caption = styles.add_style("Caption Korean", WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption = styles["Caption Korean"]
    caption.font.name = "Malgun Gothic"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
    caption.font.size = Pt(8.5)
    caption.font.italic = True
    caption.font.color.rgb = RGBColor.from_string("5E6D7C")
    caption.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    caption.paragraph_format.space_after = Pt(8)

    return doc


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run("MoonCen Architecture Review  ·  ")
    set_run_font(run, size=8, color="718096")
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = " PAGE "
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def add_header_footer(doc: Document, short_title: str) -> None:
    for sec in doc.sections:
        header = sec.header.paragraphs[0]
        header.text = short_title
        header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for run in header.runs:
            set_run_font(run, size=8, color="718096")
        p_pr = header._p.get_or_add_pPr()
        borders = OxmlElement("w:pBdr")
        bottom = OxmlElement("w:bottom")
        bottom.set(qn("w:val"), "single")
        bottom.set(qn("w:sz"), "4")
        bottom.set(qn("w:space"), "1")
        bottom.set(qn("w:color"), MID_GRAY)
        borders.append(bottom)
        p_pr.append(borders)
        add_page_number(sec.footer.paragraphs[0])


def add_toc(doc: Document) -> None:
    p = doc.add_paragraph()
    run = p.add_run("목차")
    set_run_font(run, size=18, color=NAVY, bold=True)
    p.paragraph_format.space_after = Pt(12)
    toc = doc.add_paragraph()
    begin_run = toc.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin_run._r.append(begin)
    instr_run = toc.add_run()
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    instr_run._r.append(instr)
    separate_run = toc.add_run()
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    separate_run._r.append(separate)
    placeholder_run = toc.add_run("Word에서 문서를 열면 목차가 자동으로 업데이트됩니다.")
    set_run_font(placeholder_run, size=9, color="718096")
    end_run = toc.add_run()
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    end_run._r.append(end)


def add_cover(doc: Document, title: str) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(74)
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run("MOONCEN")
    set_run_font(run, size=13, color=BLUE, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(12)
    run = p.add_run(title)
    set_run_font(run, size=29, color=NAVY, bold=True)

    p = doc.add_paragraph()
    run = p.add_run("코드·설정·SQL·배포 자산 기반 정적 분석")
    set_run_font(run, size=13, color="5E6D7C")

    doc.add_paragraph()
    table = doc.add_table(rows=3, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False
    set_repeat_table_width(table)
    info = (("작성 기준", SNAPSHOT), ("분석 revision", REVISION), ("문서 상태", "현재 워크스페이스 스냅샷 · 실제 배포 상태는 별도 검증 필요"))
    for row, values in zip(table.rows, info):
        row.cells[0].width = Cm(3.2)
        row.cells[1].width = Cm(13.8)
        for idx, value in enumerate(values):
            cell = row.cells[idx]
            set_cell_shading(cell, LIGHT_BLUE if idx == 0 else WHITE)
            set_cell_margins(cell, 120, 140, 120, 140)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(0)
            run = para.add_run(value)
            set_run_font(run, size=9.2, color=NAVY if idx == 0 else TEXT, bold=idx == 0)

    doc.add_paragraph()
    warn = doc.add_table(rows=1, cols=1)
    warn.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_repeat_table_width(warn)
    cell = warn.cell(0, 0)
    set_cell_shading(cell, "FFF4E5")
    set_cell_margins(cell, 160, 170, 160, 170)
    para = cell.paragraphs[0]
    para.paragraph_format.space_after = Pt(0)
    run = para.add_run("해석 주의  ")
    set_run_font(run, size=9.2, color=RED, bold=True)
    run = para.add_run("저장소가 선언하는 구성과 실제 운영 서버의 revision·서비스·DB role·접근 통제·백업 성공은 서로 다른 증거입니다.")
    set_run_font(run, size=9.2, color=TEXT)

    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(105)
    run = p.add_run("Prepared for MoonCen engineering and operations")
    set_run_font(run, size=9, color="718096")
    doc.add_page_break()
    add_toc(doc)
    doc.add_page_break()


INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|`[^`]+`|\[[^\]]+\]\([^\)]+\))")


def add_hyperlink(paragraph, text: str, target: str) -> None:
    # Relative repository references remain readable text; external URLs become links.
    if not re.match(r"https?://", target):
        run = paragraph.add_run(text)
        set_run_font(run, color=BLUE)
        return
    part = paragraph.part
    rel_id = part.relate_to(target, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), rel_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.extend([color, underline])
    new_run.append(r_pr)
    text_node = OxmlElement("w:t")
    text_node.text = text
    new_run.append(text_node)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def add_inline(paragraph, text: str, *, size=None, color=None, bold=False) -> None:
    cursor = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            set_run_font(run, size=size, color=color, bold=bold if bold else None)
        token = match.group(0)
        if token.startswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, size=size, color=color, bold=True)
        elif token.startswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_font(run, name="Consolas", size=(size or 9.2) - 0.2, color="7A3E00")
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
            shd = OxmlElement("w:shd")
            shd.set(qn("w:fill"), "F3F0EA")
            run._element.get_or_add_rPr().append(shd)
        else:
            link = re.match(r"\[([^\]]+)\]\(([^\)]+)\)", token)
            if link:
                add_hyperlink(paragraph, link.group(1), link.group(2))
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, size=size, color=color, bold=bold if bold else None)


def add_horizontal_rule(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p_pr = p._p.get_or_add_pPr()
    borders = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), MID_GRAY)
    borders.append(bottom)
    p_pr.append(borders)


def add_quote(doc: Document, text: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    set_repeat_table_width(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "FFF4E5" if "중요" in text or "주의" in text else LIGHT_BLUE)
    set_cell_margins(cell, 110, 150, 110, 150)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    add_inline(p, text, size=9.1)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_code_block(doc: Document, lines: list[str]) -> None:
    p = doc.add_paragraph(style="Code Block")
    p_pr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), "F4F6F8")
    p_pr.append(shd)
    borders = OxmlElement("w:pBdr")
    for edge in ("top", "left", "bottom", "right"):
        border = OxmlElement(f"w:{edge}")
        border.set(qn("w:val"), "single")
        border.set(qn("w:sz"), "4")
        border.set(qn("w:color"), MID_GRAY)
        borders.append(border)
    p_pr.append(borders)
    for idx, line in enumerate(lines):
        run = p.add_run(line)
        set_run_font(run, name="Consolas", size=8.2, color="243447")
        run._element.rPr.rFonts.set(qn("w:eastAsia"), "Malgun Gothic")
        if idx != len(lines) - 1:
            run.add_break()


def strip_inline_markup(text: str) -> str:
    text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\[([^\]]+)\]\([^\)]+\)", r"\1", text)
    return text.replace("  ", " ").strip()


def parse_table(lines: list[str], start: int) -> tuple[list[list[str]], int]:
    rows: list[list[str]] = []
    idx = start
    while idx < len(lines) and lines[idx].strip().startswith("|"):
        cells = [cell.strip() for cell in lines[idx].strip().strip("|").split("|")]
        rows.append(cells)
        idx += 1
    if len(rows) >= 2 and all(re.fullmatch(r":?-{3,}:?", cell.replace(" ", "")) for cell in rows[1]):
        rows.pop(1)
    return rows, idx


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    cols = max(len(row) for row in rows)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = True
    set_repeat_table_width(table)
    set_repeat_table_header(table.rows[0])
    font_size = 8.1 if cols >= 6 else 8.5 if cols >= 4 else 9.0
    for r_idx, values in enumerate(rows):
        for c_idx in range(cols):
            cell = table.cell(r_idx, c_idx)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell, 80, 85, 80, 85)
            set_cell_shading(cell, NAVY if r_idx == 0 else ("F7F9FB" if r_idx % 2 == 0 else WHITE))
            value = values[c_idx] if c_idx < len(values) else ""
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.0
            add_inline(p, value, size=font_size, color=WHITE if r_idx == 0 else TEXT, bold=r_idx == 0)
    doc.add_paragraph().paragraph_format.space_after = Pt(0)


def add_image(doc: Document, alt: str, rel_path: str) -> None:
    source = (ROOT / rel_path).resolve()
    if source.suffix.lower() == ".svg":
        image_path = source.with_suffix(".png")
    else:
        image_path = source
    if not image_path.exists():
        raise FileNotFoundError(f"Rendered image missing: {image_path}")
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run()
    run.add_picture(str(image_path), width=Inches(6.65))
    caption = doc.add_paragraph(style="Caption Korean")
    caption.add_run(alt)


def markdown_to_docx(source: Path) -> Path:
    lines = source.read_text(encoding="utf-8").splitlines()
    if not lines or not lines[0].startswith("# "):
        raise ValueError(f"Document must begin with H1: {source}")
    title = lines[0][2:].strip()
    doc = setup_document()
    doc.core_properties.title = title
    doc.core_properties.subject = "MoonCen architecture, improvements, and operations review"
    doc.core_properties.author = "OpenAI Codex — repository-based analysis"
    doc.core_properties.keywords = "MoonCen, architecture, operations, crawler, FastAPI, PostgreSQL"
    add_header_footer(doc, title)
    add_cover(doc, title)

    idx = 1
    in_code = False
    code_lines: list[str] = []
    while idx < len(lines):
        raw = lines[idx]
        stripped = raw.strip()

        if stripped.startswith("```"):
            if in_code:
                add_code_block(doc, code_lines)
                code_lines = []
                in_code = False
            else:
                in_code = True
            idx += 1
            continue
        if in_code:
            code_lines.append(raw)
            idx += 1
            continue
        if not stripped:
            idx += 1
            continue
        if stripped == "---":
            add_horizontal_rule(doc)
            idx += 1
            continue

        image_match = re.fullmatch(r"!\[([^\]]*)\]\(([^\)]+)\)", stripped)
        if image_match:
            add_image(doc, image_match.group(1), image_match.group(2))
            idx += 1
            continue

        if stripped.startswith("|") and idx + 1 < len(lines) and lines[idx + 1].strip().startswith("|"):
            rows, idx = parse_table(lines, idx)
            add_table(doc, rows)
            continue

        heading = re.match(r"^(#{2,4})\s+(.+)$", stripped)
        if heading:
            level = len(heading.group(1)) - 1
            p = doc.add_paragraph(style=f"Heading {level}")
            add_inline(p, heading.group(2))
            set_keep_with_next(p)
            idx += 1
            continue

        if stripped.startswith("> "):
            quote_parts = [stripped[2:]]
            idx += 1
            while idx < len(lines) and lines[idx].strip().startswith("> "):
                quote_parts.append(lines[idx].strip()[2:])
                idx += 1
            add_quote(doc, " ".join(quote_parts))
            continue

        list_match = re.match(r"^(\s*)[-*]\s+(.*)$", raw)
        number_match = re.match(r"^(\s*)\d+\.\s+(.*)$", raw)
        if list_match or number_match:
            match = list_match or number_match
            assert match is not None
            depth = min(len(match.group(1)) // 2, 2)
            style_root = "List Bullet" if list_match else "List Number"
            style = style_root if depth == 0 else f"{style_root} {depth + 1}"
            if style not in doc.styles:
                style = style_root
            p = doc.add_paragraph(style=style)
            content = match.group(2)
            if content.startswith("[ ] "):
                content = "☐ " + content[4:]
            elif content.startswith("[x] ") or content.startswith("[X] "):
                content = "☑ " + content[4:]
            add_inline(p, content)
            idx += 1
            continue

        paragraph_parts = [stripped]
        idx += 1
        while idx < len(lines):
            next_raw = lines[idx]
            next_stripped = next_raw.strip()
            if not next_stripped:
                idx += 1
                break
            if (
                next_stripped.startswith(("#", "```", "> ", "|", "!["))
                or next_stripped == "---"
                or re.match(r"^\s*[-*]\s+", next_raw)
                or re.match(r"^\s*\d+\.\s+", next_raw)
            ):
                break
            paragraph_parts.append(next_stripped)
            idx += 1
        p = doc.add_paragraph()
        add_inline(p, " ".join(paragraph_parts))

    if in_code:
        add_code_block(doc, code_lines)

    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    output = source.with_suffix(".docx")
    doc.save(output)
    return output


FONT_REGULAR = "/usr/share/fonts/truetype/dejavu/DejaVuSans.ttf"
FONT_BOLD = "/usr/share/fonts/truetype/dejavu/DejaVuSans-Bold.ttf"


def image_font(size: int, bold: bool = False):
    path = FONT_BOLD if bold else FONT_REGULAR
    return ImageFont.truetype(path, size=size)


def hex_color(value: str) -> str:
    return "#" + value.lstrip("#")


def draw_centered(draw: ImageDraw.ImageDraw, xy, text: str, size: int, color: str, bold=False) -> None:
    draw.text(xy, text, font=image_font(size, bold), fill=hex_color(color), anchor="mm")


def wrap_pixels(draw: ImageDraw.ImageDraw, text: str, font, max_width: int) -> list[str]:
    words = text.split()
    if not words:
        return [""]
    lines: list[str] = []
    current = words[0]
    for word in words[1:]:
        candidate = current + " " + word
        if draw.textbbox((0, 0), candidate, font=font)[2] <= max_width:
            current = candidate
        else:
            lines.append(current)
            current = word
    lines.append(current)
    return lines


def draw_box(
    draw: ImageDraw.ImageDraw,
    bounds,
    title: str,
    lines: list[str],
    *,
    fill: str,
    outline: str,
    title_size=21,
    line_size=15,
) -> None:
    x1, y1, x2, y2 = bounds
    draw.rounded_rectangle(bounds, radius=17, fill=hex_color(fill), outline=hex_color(outline), width=3)
    title_font = image_font(title_size, True)
    line_font = image_font(line_size)
    text_lines: list[tuple[str, object, str]] = [(title, title_font, NAVY)]
    for line in lines:
        for wrapped in wrap_pixels(draw, line, line_font, max(60, x2 - x1 - 28)):
            text_lines.append((wrapped, line_font, TEXT if not line.startswith("[") else "5E6D7C"))
    heights = [draw.textbbox((0, 0), value, font=font)[3] + (8 if idx == 0 else 5) for idx, (value, font, _color) in enumerate(text_lines)]
    total = sum(heights)
    y = y1 + (y2 - y1 - total) / 2
    for idx, ((value, font, color), height) in enumerate(zip(text_lines, heights)):
        draw.text(((x1 + x2) / 2, y + height / 2), value, font=font, fill=hex_color(color), anchor="mm")
        y += height


def draw_arrow(draw: ImageDraw.ImageDraw, points, *, color="476782", width=4, dashed=False) -> None:
    rgb = hex_color(color)
    if dashed:
        for start, end in zip(points, points[1:]):
            x1, y1 = start
            x2, y2 = end
            length = math.hypot(x2 - x1, y2 - y1)
            if length == 0:
                continue
            ux, uy = (x2 - x1) / length, (y2 - y1) / length
            pos = 0.0
            while pos < length - 15:
                stop = min(pos + 12, length - 15)
                draw.line((x1 + ux * pos, y1 + uy * pos, x1 + ux * stop, y1 + uy * stop), fill=rgb, width=width)
                pos += 21
    else:
        draw.line(points, fill=rgb, width=width, joint="curve")
    x1, y1 = points[-2]
    x2, y2 = points[-1]
    angle = math.atan2(y2 - y1, x2 - x1)
    head = 16
    wing = 8
    left = (x2 - head * math.cos(angle) + wing * math.sin(angle), y2 - head * math.sin(angle) - wing * math.cos(angle))
    right = (x2 - head * math.cos(angle) - wing * math.sin(angle), y2 - head * math.sin(angle) + wing * math.cos(angle))
    draw.polygon([(x2, y2), left, right], fill=rgb)


def base_canvas(height: int, title: str, subtitle: str):
    image = Image.new("RGB", (1600, height), "#F6F8FB")
    draw = ImageDraw.Draw(image)
    draw.text((65, 42), title, font=image_font(34, True), fill="#132238")
    draw.text((65, 87), subtitle, font=image_font(17), fill="#56657A")
    return image, draw


def render_current_architecture(path: Path) -> None:
    image, draw = base_canvas(
        980,
        "MoonCen - Current Declared Architecture",
        "Repository snapshot 2026-08-19 | crawlerMode=legacy | live deployment requires verification",
    )
    draw_box(draw, (65, 140, 310, 270), "Public Clients", ["Web / SEO / Mobile", "[Browser and Expo read-only]"], fill="E8F1FF", outline="5F8ECB")
    draw_box(draw, (390, 120, 720, 290), "Cloudflare + Nginx", ["TLS, edge controls, rate limits", "Public host blocks /api/ops", "SPA / SEO reverse proxy"], fill="FFFFFF", outline="8CA3BD")
    draw_box(draw, (805, 115, 1135, 295), "cloud - Application", ["frontend2 and FastAPI", "Auth, public API and SEO", "Ops API behind a separate boundary"], fill="EAF7F2", outline="4F9B7A")
    draw_box(draw, (1245, 120, 1535, 290), "Primary PostgreSQL", ["PostGIS, courses and branches", "Users and ops metadata", "[No standby in current manifest]"], fill="FFF4DF", outline="D29B3D")
    draw_arrow(draw, [(310, 205), (390, 205)])
    draw_arrow(draw, [(720, 205), (805, 205)])
    draw_arrow(draw, [(1135, 205), (1245, 205)])

    draw_box(draw, (65, 385, 375, 565), "gen1crawler", ["Legacy scheduler / orchestrator", "Provider crawlers and maintenance", "[Declared crawler owner]"], fill="EDE9FF", outline="7868B4")
    draw_box(draw, (475, 360, 820, 590), "gen1db - Staging / Control", ["Batches and staged rows", "Validation evidence and snapshots", "Control-plane metadata", "[Distributed capability inactive]"], fill="FFF4DF", outline="D29B3D")
    draw_box(draw, (920, 385, 1230, 565), "Promotion Gate", ["validate -> dry-run fingerprint", "review -> pinned transaction", "Safe close-missing and rollback"], fill="FDEEED", outline="C55F55")
    draw_box(draw, (1290, 385, 1535, 565), "External Dependencies", ["Source sites, maps, OAuth", "SMTP, Cloudflare, AI", "[Availability outside repository]"], fill="FFFFFF", outline="8CA3BD", line_size=14)
    draw_arrow(draw, [(375, 475), (475, 475)])
    draw_arrow(draw, [(820, 475), (920, 475)])
    draw_arrow(draw, [(1075, 385), (1075, 330), (1380, 330), (1380, 290)])
    draw_arrow(draw, [(1290, 500), (1230, 500)], color="A04B40", dashed=True)

    draw_box(draw, (65, 710, 350, 865), "Ops Browser", ["Separate ops origin", "viewer / operator / admin", "Polling and SSE"], fill="E8F1FF", outline="5F8ECB")
    draw_box(draw, (455, 680, 810, 895), "Operations Control", ["ops-console and FastAPI Ops", "ops_agent jobs and deployments", "Audit, quality and crawler control", "[Access and MFA must be verified]"], fill="EAF7F2", outline="4F9B7A")
    draw_box(draw, (920, 690, 1230, 885), "bot - Observability", ["Prometheus and Grafana", "Node/textfile metrics and alerts", "[Some targets remain pending]"], fill="EDF7FB", outline="5791AA")
    draw_box(draw, (1300, 710, 1535, 865), "Backup Target", ["Encrypted NAS copy", "Daily job and monthly drill", "[Success evidence required]"], fill="F2F2F2", outline="7D8794", line_size=14)
    draw_arrow(draw, [(350, 785), (455, 785)])
    draw_arrow(draw, [(810, 785), (920, 785)])
    draw_arrow(draw, [(1230, 785), (1300, 785)])
    draw_arrow(draw, [(630, 680), (630, 590)])
    draw_arrow(draw, [(1075, 690), (1075, 565)])
    draw.text((65, 943), "Solid arrows: declared data/control path. Dashed arrow: external boundary. Not proof of live deployment.", font=image_font(14), fill="#56657A")
    image.save(path, "PNG", optimize=True)


def render_crawler_flow(path: Path) -> None:
    image, draw = base_canvas(
        990,
        "Crawler Data Flow - Current and Target",
        "Promotion safeguards are implemented; distributed execution remains explicitly gated",
    )
    draw.rounded_rectangle((45, 120, 1555, 455), radius=24, fill="#FFFFFF", outline="#7AA285", width=4)
    draw.text((80, 150), "A. CURRENT DECLARED MODE: LEGACY", font=image_font(25, True), fill="#132238")
    current = [
        ((80, 220, 320, 365), "Provider Sources", ["Public / partner sites", "42 logical groups"], "EDE9FF", "7868B4"),
        ((390, 200, 650, 385), "gen1crawler", ["run_crawlers.py", "fixed-argv subprocess", "lock, timeout, progress"], "E8F1FF", "5F8ECB"),
        ((725, 200, 985, 385), "Staging Batch", ["branches and courses", "batch/provider ownership", "validation snapshots"], "FFF4DF", "D29B3D"),
        ((1060, 190, 1335, 395), "Promotion Controls", ["completeness and quality", "reviewed fingerprint", "transactional apply", "safe close-missing"], "FDEEED", "C55F55"),
        ((1405, 220, 1520, 365), "Primary", ["Postgres", "PostGIS"], "EAF7F2", "4F9B7A"),
    ]
    for bounds, title, lines, fill, outline in current:
        draw_box(draw, bounds, title, lines, fill=fill, outline=outline, title_size=18, line_size=14)
    for p1, p2 in [((320, 292), (390, 292)), ((650, 292), (725, 292)), ((985, 292), (1060, 292)), ((1335, 292), (1405, 292))]:
        draw_arrow(draw, [p1, p2])
    draw.text((80, 420), "A batch is not production data until evidence and reviewed fingerprint match at apply time.", font=image_font(14), fill="#56657A")

    draw.rounded_rectangle((45, 500, 1555, 925), radius=24, fill="#FFFAFA", outline="#BD4C43", width=4)
    draw.text((80, 530), "B. TARGET CAPABILITY: DISTRIBUTED - NOT READY / ACTIVATION BLOCKED", font=image_font(24, True), fill="#A43E36")
    target = [
        ((80, 630, 325, 795), "Control Plane", ["artifact and desired state", "task / attempt / lease", "[Declared on gen1db]"], "FFF4DF", "D29B3D"),
        ((405, 595, 685, 830), "Canary-first Workers", ["wtr-linux: disabled", "gen1crawler: disabled", "434 concrete tasks", "resource and ownership limits"], "EDE9FF", "7868B4"),
        ((765, 605, 1025, 820), "Immutable Evidence", ["attempt observations", "receipts and snapshots", "finalizer / publisher", "fenced ownership"], "E8F1FF", "5F8ECB"),
        ((1100, 605, 1360, 820), "Same Promotion Gate", ["complete attempts", "reviewed fingerprint", "exact provider apply", "no primary bypass"], "FDEEED", "C55F55"),
        ((1420, 630, 1520, 795), "Primary", ["DB", "after approval"], "EAF7F2", "4F9B7A"),
    ]
    for bounds, title, lines, fill, outline in target:
        draw_box(draw, bounds, title, lines, fill=fill, outline=outline, title_size=18, line_size=14)
    for p1, p2 in [((325, 712), (405, 712)), ((685, 712), (765, 712)), ((1025, 712), (1100, 712)), ((1360, 712), (1420, 712))]:
        draw_arrow(draw, [p1, p2], color="BD4C43", dashed=True)
    draw.rounded_rectangle((340, 855, 1390, 900), radius=12, fill="#BD4C43")
    draw_centered(draw, (865, 877), "GO requires backup/restore, signed artifact, bootstrap, canary, monitoring and atomic cutover evidence", 17, WHITE, True)
    draw.text((65, 955), "The target path is code/design capability, not evidence that distributed crawling runs in production.", font=image_font(14), fill="#56657A")
    image.save(path, "PNG", optimize=True)


def render_operations_loop(path: Path) -> None:
    image, draw = base_canvas(
        920,
        "Recommended Production Operations Loop",
        "Evidence-driven release, monitoring, recovery and continuous improvement",
    )
    boxes = [
        ((95, 205, 390, 370), "1. Plan & Approve", ["owner, risk, rollback, window", "topology and data impact"], "E8F1FF", "5F8ECB"),
        ((500, 125, 810, 295), "2. Build & Verify", ["clean commit and CI", "signed hash, SBOM, migration proof"], "EDE9FF", "7868B4"),
        ((945, 125, 1255, 295), "3. Release Safely", ["backup and preflight", "canary, health, smoke, rollback"], "FFF4DF", "D29B3D"),
        ((1225, 355, 1525, 530), "4. Observe", ["availability, latency, errors", "crawler, DB and client health"], "EAF7F2", "4F9B7A"),
        ((945, 625, 1255, 800), "5. Respond & Recover", ["triage, contain, preserve evidence", "rollback or restore, communicate"], "FDEEED", "C55F55"),
        ((500, 625, 810, 800), "6. Learn & Improve", ["review and owned actions", "runbook, test and alert tuning"], "EDF7FB", "5791AA"),
        ((95, 535, 390, 710), "7. Reconcile Evidence", ["deployed hash and inventory", "backup, drill and access proof"], "F2F2F2", "7D8794"),
    ]
    for bounds, title, lines, fill, outline in boxes:
        draw_box(draw, bounds, title, lines, fill=fill, outline=outline, title_size=20, line_size=14)
    arrows = [
        [(390, 250), (500, 210)],
        [(810, 210), (945, 210)],
        [(1255, 245), (1375, 355)],
        [(1375, 530), (1255, 680)],
        [(945, 710), (810, 710)],
        [(500, 710), (390, 650)],
        [(240, 535), (240, 370)],
    ]
    for points in arrows:
        draw_arrow(draw, points)
    draw.ellipse((680, 350, 920, 590), fill="#132238")
    draw_centered(draw, (800, 430), "OPERATE FROM", 24, WHITE, True)
    draw_centered(draw, (800, 470), "VERIFIABLE", 24, WHITE, True)
    draw_centered(draw, (800, 510), "EVIDENCE", 24, WHITE, True)
    draw.text((65, 875), "Repository controls are a strong base; live status, access, backup success and deployment identity still need proof.", font=image_font(14), fill="#56657A")
    image.save(path, "PNG", optimize=True)


def render_diagram_with_pillow(stem: str, path: Path) -> None:
    renderers = {
        "01-current-architecture": render_current_architecture,
        "02-crawler-flow": render_crawler_flow,
        "03-operations-loop": render_operations_loop,
    }
    renderer = renderers.get(stem)
    if renderer is None:
        raise ValueError(f"No Pillow renderer registered for {stem}")
    renderer(path)


def render_assets() -> list[Path]:
    outputs: list[Path] = []
    for svg in sorted(ASSETS.glob("*.svg")):
        png = svg.with_suffix(".png")
        if cairosvg is not None:
            cairosvg.svg2png(url=str(svg), write_to=str(png), output_width=1600)
        else:
            render_diagram_with_pillow(svg.stem, png)
        outputs.append(png)
    return outputs


def validate_sources() -> None:
    missing = [path for path in SOURCE_FILES if not path.exists()]
    if missing:
        raise FileNotFoundError("Missing source files: " + ", ".join(str(path) for path in missing))
    for source in SOURCE_FILES:
        text = source.read_text(encoding="utf-8")
        for rel in re.findall(r"!\[[^\]]*\]\(([^\)]+)\)", text):
            if not (ROOT / rel).exists():
                raise FileNotFoundError(f"Missing image reference in {source.name}: {rel}")


def main() -> int:
    validate_sources()
    rendered = render_assets()
    outputs = [markdown_to_docx(source) for source in SOURCE_FILES]
    print("Rendered PNG:")
    for path in rendered:
        print(f"  {path.name} ({path.stat().st_size:,} bytes)")
    print("Built DOCX:")
    for path in outputs:
        print(f"  {path.name} ({path.stat().st_size:,} bytes)")
    return 0


if __name__ == "__main__":
    sys.exit(main())
